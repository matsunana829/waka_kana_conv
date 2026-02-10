from __future__ import annotations

import io
from typing import Iterable, Tuple
import re
import xml.etree.ElementTree as ET

import pandas as pd
from docx import Document


def read_text_bytes(data: bytes) -> str:
    # 文字コードを推定してテキストに変換
    for enc in ("utf-8-sig", "utf-8", "cp932"):
        try:
            return data.decode(enc)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def write_csv(df: pd.DataFrame) -> bytes:
    # CSVをUTF-8(BOM付き)で出力
    buffer = io.StringIO()
    df.to_csv(buffer, index=False)
    return buffer.getvalue().encode("utf-8-sig")


def convert_docx_bytes(
    data: bytes,
    convert_func,
) -> Tuple[bytes, bytes]:
    # docxの段落テキストを変換し、txtとdocxの両方を返す
    doc = Document(io.BytesIO(data))
    out_doc = Document()
    text_lines = []
    for para in doc.paragraphs:
        converted = convert_func(para.text)
        out_doc.add_paragraph(converted)
        text_lines.append(converted)

    text_bytes = ("\n".join(text_lines)).encode("utf-8-sig")
    out_stream = io.BytesIO()
    out_doc.save(out_stream)
    return text_bytes, out_stream.getvalue()


def convert_csv_bytes(
    data: bytes,
    text_column: str,
    convert_func,
) -> pd.DataFrame:
    # CSVを読み込み、指定列だけ変換してDataFrameで返す
    df = None
    for enc in ("utf-8-sig", "utf-8", "cp932"):
        try:
            df = pd.read_csv(io.BytesIO(data), encoding=enc)
            break
        except Exception:
            df = None
    if df is None:
        df = pd.read_csv(io.BytesIO(data), encoding="utf-8", errors="replace")

    if text_column not in df.columns:
        raise KeyError(f"Column not found: {text_column}")

    df[text_column] = df[text_column].astype(str).map(convert_func)
    return df


def _local_name(tag: str) -> str:
    # 名前空間付きタグからローカル名だけを取り出す
    if "}" in tag:
        return tag.split("}", 1)[1]
    return tag


def _register_namespaces(xml_text: str) -> None:
    # 元XMLの名前空間宣言を登録して、出力時のns0化を防ぐ
    for match in re.finditer(r'xmlns(?::([\\w.-]+))?="([^"]+)"', xml_text):
        prefix, uri = match.group(1), match.group(2)
        if prefix is None:
            ET.register_namespace("", uri)
        else:
            ET.register_namespace(prefix, uri)


def _is_hiragana(ch: str) -> bool:
    return 0x3041 <= ord(ch) <= 0x3096


def _dakuten_map() -> dict[str, str]:
    return {
        "か": "が",
        "き": "ぎ",
        "く": "ぐ",
        "け": "げ",
        "こ": "ご",
        "さ": "ざ",
        "し": "じ",
        "す": "ず",
        "せ": "ぜ",
        "そ": "ぞ",
        "た": "だ",
        "ち": "ぢ",
        "つ": "づ",
        "て": "で",
        "と": "ど",
        "は": "ば",
        "ひ": "び",
        "ふ": "ぶ",
        "へ": "べ",
        "ほ": "ぼ",
        "う": "ゔ",
    }


def _pre_expand_odoriji_in_element(elem: ET.Element) -> None:
    # <l>内の全テキストをタグ無視で走査し、踊り字を事前展開する
    dakuten = _dakuten_map()
    prev = ""

    def process_text(text: str) -> str:
        nonlocal prev
        if not text:
            return text
        out = []
        for ch in text:
            if ch in ("ゝ", "ヽ"):
                if _is_hiragana(prev):
                    out.append(prev)
                else:
                    out.append(ch)
                continue
            if ch in ("ゞ", "ヾ"):
                if _is_hiragana(prev):
                    replaced = dakuten.get(prev, prev)
                    out.append(replaced)
                    prev = replaced
                else:
                    out.append(ch)
                continue
            out.append(ch)
            prev = ch
        return "".join(out)

    def walk(node: ET.Element) -> None:
        if node.text:
            node.text = process_text(node.text)
        for child in list(node):
            walk(child)
            if child.tail:
                child.tail = process_text(child.tail)

    walk(elem)


def _convert_text_in_element(elem: ET.Element, convert_func) -> None:
    # 指定要素配下のテキストとtailを再帰的に変換
    if elem.text:
        elem.text = convert_func(elem.text)
    for child in list(elem):
        _convert_text_in_element(child, convert_func)
        if child.tail:
            child.tail = convert_func(child.tail)


def convert_xml_bytes(
    data: bytes,
    text_tag: str,
    convert_func,
    pre_expand_odoriji: bool = False,
) -> bytes:
    # XMLを読み込み、指定タグ配下のテキストのみを変換する
    try:
        xml_text = data.decode("utf-8")
    except UnicodeDecodeError:
        xml_text = data.decode("utf-8", errors="replace")
    _register_namespaces(xml_text)

    tree = ET.ElementTree(ET.fromstring(data))
    root = tree.getroot()

    for elem in root.iter():
        if _local_name(elem.tag) != text_tag:
            continue
        if pre_expand_odoriji:
            _pre_expand_odoriji_in_element(elem)
        _convert_text_in_element(elem, convert_func)

    out = io.BytesIO()
    tree.write(out, encoding="utf-8", xml_declaration=True)
    return out.getvalue()


def texts_to_txt_bytes(texts: Iterable[str]) -> bytes:
    # テキスト配列を1つのtxtにまとめる
    return ("\n".join(texts)).encode("utf-8-sig")


def texts_to_csv_bytes(texts: Iterable[str]) -> bytes:
    # テキスト配列を1列CSVに変換
    df = pd.DataFrame({"text": list(texts)})
    return write_csv(df)


def texts_to_xml_bytes(texts: Iterable[str]) -> bytes:
    # テキスト配列を <root><text>...</text></root> 形式で出力
    root = ET.Element("root")
    for t in texts:
        elem = ET.SubElement(root, "text")
        elem.text = t
    out = io.BytesIO()
    ET.ElementTree(root).write(out, encoding="utf-8", xml_declaration=True)
    return out.getvalue()
